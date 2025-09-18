import whisper
from keybert import KeyBERT
from transformers import pipeline
import nltk
from nltk.tokenize import sent_tokenize
from docx import Document
from gpt4all import GPT4All  # Requires 'gpt4all' package and downloaded model
from datetime import datetime

# Download NLTK data if not already present
try:
    nltk.data.find('tokenizers/punkt')
except nltk.downloader.DownloadError:
    nltk.download('punkt')

def main():
    # --- Configuration ---
    # Replace 'your_audio_file.mp3' with the name of your meeting audio file
    audio_file = "audio_files/small_meeting.mp3"  

    # Replace 'mistral-7b-instruct-v0.1.Q4_0.gguf' with your downloaded GPT4All model filename
    gpt4all_model_path = "C:/Users/abhis/AppData/Local/nomic.ai/GPT4All/Phi-3-mini-4k-instruct.Q4_0.gguf"

    # --- Load Models ---
    print("Loading Whisper model...")
    whisper_model = whisper.load_model("medium")

    print("Loading KeyBERT model...")
    kw_model = KeyBERT(model='all-MiniLM-L6-v2')

    print("Loading summarizer (BART)...")
    summarizer = pipeline("summarization", model="facebook/bart-large-cnn")

    print("Loading GPT4All model for action extraction...")
    gpt4all_model = GPT4All(gpt4all_model_path)

    # --- Transcription ---
    print(f"Transcribing audio from {audio_file} ...")
    transcription_result = whisper_model.transcribe(audio_file)
    transcript = transcription_result['text']
    print("Transcription complete.\n")
    print("\n---------------------------\n")

    # --- Keyword Extraction ---
    keywords = kw_model.extract_keywords(transcript, keyphrase_ngram_range=(1, 3), top_n=30)
    
    # Post-process to filter out similar keywords (optional but recommended)
    from keybert.backend import SentenceTransformerBackend
    kw_model_st = KeyBERT(model=SentenceTransformerBackend('all-MiniLM-L6-v2'))
    filtered_keywords = kw_model_st.extract_keywords(transcript, keyphrase_ngram_range=(1, 3), stop_words='english', top_n=30, use_mmr=True, diversity=0.7)
    keywords = filtered_keywords
    print("Keywords extracted.")

    # --- Sentence Tokenization ---
    sentences = sent_tokenize(transcript)

    # --- Select sentences containing keywords ---
    keyword_words = set()
    for phrase, score in keywords:
        for w in phrase.lower().split():
            keyword_words.add(w)

    candidate_sentences = [s for s in sentences if any(k in s.lower() for k in keyword_words)]

# --- Action Extraction via GPT4All ---
    extracted_actions = []

    print("Extracting actions using GPT4All...")

    # We need to chunk sentences to provide the LLM with enough context
    chunk_size = 4
    candidate_sentences_chunks = [candidate_sentences[i:i + chunk_size] for i in range(0, len(candidate_sentences), chunk_size)]

    for chunk in candidate_sentences_chunks:
        chunk_text = " ".join(chunk)
        prompt = (
            "<|user|>\n"
            "You are a helpful meeting assistant. Your task is to extract only the specific action items, decisions, or assigned tasks from the text below.\n"
            "Identify the owner of the task if a person's name or title is mentioned. The owner of the task is the person responsible for completing the task.\n"
            "If no clear action items are found, respond with 'None'.\n"
            "Format the output as a numbered list. Be extremely concise and do not add any extra commentary or conversational text.\n"
            "Text:\n"
            f"'''{chunk_text}'''\n"
            "<|assistant|>\n"
        )
        
        with gpt4all_model.chat_session():
            output = gpt4all_model.generate(prompt, max_tokens=150)
            action_text = output.strip()
            if action_text.lower() != 'none' and len(action_text) > 5:
                extracted_actions.append(action_text)

    # Remove duplicates
    extracted_actions = list(dict.fromkeys(extracted_actions))

    # --- Summarize Transcript ---
    print("Summarizing transcript...")
    max_chunk = 500  # approx character chunk size to avoid model limits
    chunks = [transcript[i:i + max_chunk] for i in range(0, len(transcript), max_chunk)]
    summary_parts = []

    for chunk in chunks:
        summary_result = summarizer(chunk, max_length=130, min_length=30, do_sample=False)
        summary_parts.append(summary_result[0]['summary_text'])
    summary = " ".join(summary_parts)

    # --- Save Results to DOCX File ---
    print("Saving results to a new DOCX file...")
    doc = Document()
    doc.add_heading("Automated Meeting Notes", level=0)
    doc.add_heading("Transcript:", level=1)
    doc.add_paragraph(transcript)
    doc.add_heading("Keywords:", level=1)
    for kw, score in keywords:
        doc.add_paragraph(f"{kw} (score: {score:.3f})")
    doc.add_heading("Extracted Actions / Decisions:", level=1)
    if extracted_actions:
        for act in extracted_actions:
            doc.add_paragraph(act, style='List Bullet')
    else:
        doc.add_paragraph("No action items detected.")
    doc.add_heading("Summary:", level=1)
    doc.add_paragraph(summary)
    
    # Get the current date and time to create a unique filename
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    output_filename = f"meeting_notes_{timestamp}.docx"
    
    doc.save(output_filename)

    print(f"Done! Check '{output_filename}' for outputs.")

if __name__ == "__main__":
    main()